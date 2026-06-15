#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
简历解析模块
功能：
1. 读取PDF文件
2. 提取文本内容
3. 调用LLM提取结构化信息
4. 调用技能标签打分功能
"""
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

import pandas as pd
import PyPDF2
import requests
import time

try:
    from tag_rate import (
        APIKeyManager,
        COMMON_SCORING_RULES_V4,
        load_api_keys,
        parse_llm_response,
    )
except ImportError:
    logging.error("无法导入 tag_rate 模块，请确保 tag_rate.py 在同一目录")
    raise

from llm_client import LLMClient, resolve_secret_value

# 配置
ROOT_DIR = Path(__file__).resolve().parent
DEFAULT_LABELS_FILE = ROOT_DIR / "all_labels.csv"
DEFAULT_TAXONOMY_FILE = ROOT_DIR / "tech_taxonomy.json"
DEFAULT_API_KEY_FILE = ROOT_DIR / "API_key-openai.md"

_TAG_SCORE_RE = re.compile(r"([^:：\s]+?)\s*[:：]\s*([1-5])(?=\s|$)")


class ResumeParser:
    """简历解析器"""
    
    def __init__(self):
        import os
        if DEFAULT_API_KEY_FILE.exists():
            api_keys = [resolve_secret_value(key) for key in load_api_keys(DEFAULT_API_KEY_FILE)]
            self.api_key_manager = APIKeyManager(api_keys)
        else:
            env_key = (
                os.getenv("OPENROUTER_API_KEY")
                or os.getenv("OPENAI_API_KEY")
                or os.getenv("DEEPSEEK_API_KEY")
                or os.getenv("API_KEY")
                or os.getenv("LLM_API_KEY")
                or ""
            )
            self.api_key_manager = APIKeyManager([resolve_secret_value(env_key)]) if env_key else None
        self.llm = LLMClient(self.api_key_manager)
        
        # 加载技能标签库与岗位族(level_3rd -> tags)
        self.all_tags = self._load_tags()
        self.level3_list, self.level3_to_tags = self._load_level3_and_tags()
        logging.info(f"已加载 {len(self.all_tags)} 个技能标签，{len(self.level3_list)} 个岗位类别(level_3rd)")
    
    def _load_taxonomy(self):
        """加载 tech_taxonomy.json，返回 (level3_list, level3_to_tags)。
        在 all_labels.csv 缺失时作为标签库的兜底来源。
        将 level2_roles 的 name 当作 level_3rd，其 keywords 当作 tags。"""
        if not DEFAULT_TAXONOMY_FILE.exists():
            return [], {}
        try:
            data = json.loads(DEFAULT_TAXONOMY_FILE.read_text(encoding="utf-8"))
        except Exception as e:
            logging.warning(f"无法解析 tech_taxonomy.json: {e}")
            return [], {}
        level3_list: List[str] = []
        level3_to_tags: Dict[str, List[str]] = {}
        for cat in data.get("level1_categories", []):
            for role in cat.get("level2_roles", []):
                name = str(role.get("name", "")).strip()
                if not name:
                    continue
                kws = [str(k).strip() for k in role.get("keywords", []) if str(k).strip()]
                # role 名称本身也作为一个可匹配标签
                tags = [name] + kws
                if name not in level3_to_tags:
                    level3_to_tags[name] = []
                    level3_list.append(name)
                level3_to_tags[name].extend(tags)
        return level3_list, level3_to_tags

    def _load_tags(self) -> Set[str]:
        """加载所有可用的技能标签。优先 all_labels.csv，缺失时回退 tech_taxonomy.json。"""
        if not DEFAULT_LABELS_FILE.exists():
            logging.warning(f"技能标签文件不存在: {DEFAULT_LABELS_FILE}，回退 tech_taxonomy.json")
            _, level3_to_tags = self._load_taxonomy()
            tags: Set[str] = set()
            for tag_list in level3_to_tags.values():
                for t in tag_list:
                    if t:
                        tags.add(t)
            if not tags:
                logging.warning("tech_taxonomy.json 也不可用，技能标签库为空")
            return tags

        tags = set()
        df = pd.read_csv(DEFAULT_LABELS_FILE).fillna("")
        for _, row in df.iterrows():
            tags_raw = str(row.get("tags", "")).split("|_|")
            for tag in tags_raw:
                tag = tag.strip()
                if tag:
                    tags.add(tag)

        return tags

    def _load_level3_and_tags(self) -> Tuple[List[str], Dict[str, List[str]]]:
        """加载 level_3rd 列表及其对应的 tags。优先 all_labels.csv，缺失时回退 tech_taxonomy.json。"""
        level3_to_tags: Dict[str, List[str]] = {}
        level3_list: List[str] = []
        if not DEFAULT_LABELS_FILE.exists():
            level3_list, level3_to_tags = self._load_taxonomy()
            # 去重保持顺序
            for lv3, tags in level3_to_tags.items():
                seen = set()
                deduped = []
                for t in tags:
                    if t not in seen:
                        seen.add(t)
                        deduped.append(t)
                level3_to_tags[lv3] = deduped
            return level3_list, level3_to_tags
        df = pd.read_csv(DEFAULT_LABELS_FILE).fillna("")
        for _, row in df.iterrows():
            lv3 = str(row.get("level_3rd", "")).strip()
            if not lv3:
                continue
            if lv3 not in level3_to_tags:
                level3_to_tags[lv3] = []
                level3_list.append(lv3)
            tags_raw = str(row.get("tags", "")).split("|_|")
            for tag in tags_raw:
                tag = tag.strip()
                if tag:
                    level3_to_tags[lv3].append(tag)
        # 去重保持顺序
        for lv3, tags in level3_to_tags.items():
            seen = set()
            deduped = []
            for t in tags:
                if t not in seen:
                    seen.add(t)
                    deduped.append(t)
            level3_to_tags[lv3] = deduped
        return level3_list, level3_to_tags
    
    def extract_text(self, file_path: str) -> str:
        """按文件类型分派文本提取：PDF / Word(.docx) / 纯文本。

        历史问题：原先无论扩展名一律走 PDF 解析器，导致 .docx 报
        "EOF marker not found"（PyPDF2 找不到 PDF 的 %%EOF 标记）。
        现按扩展名分派，且 PDF 优先用更稳健的 pdfplumber。
        """
        ext = Path(file_path).suffix.lower().lstrip(".")
        if ext == "pdf":
            return self.extract_text_from_pdf(file_path)
        if ext == "docx":
            return self.extract_text_from_docx(file_path)
        if ext == "doc":
            # 旧版 .doc 二进制格式，python-docx 不支持，给出可操作的提示
            raise ValueError(
                "暂不支持旧版 .doc 格式，请将简历另存为 .docx 或 PDF 后重新上传"
            )
        if ext in ("txt", "md"):
            return Path(file_path).read_text(encoding="utf-8", errors="ignore").strip()
        raise ValueError(f"不支持的简历格式: .{ext}（支持 PDF / .docx）")

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """从PDF文件中提取文本。优先 pdfplumber（对真实简历更稳健），
        失败再回退 PyPDF2，两者都失败才抛错。"""
        # 1) 首选 pdfplumber
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or "")
            text = "\n".join(parts).strip()
            if text:
                return text
            logging.warning("pdfplumber 未提取到文本，尝试 PyPDF2 回退")
        except Exception as e:
            logging.warning(f"pdfplumber 提取失败，回退 PyPDF2: {e}")

        # 2) 回退 PyPDF2
        try:
            text = ""
            with open(pdf_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += (page.extract_text() or "") + "\n"
            return text.strip()
        except Exception as e:
            logging.error(f"PDF文本提取失败: {e}")
            raise ValueError(f"PDF 解析失败，文件可能已损坏或为扫描件: {e}")

    def extract_text_from_docx(self, docx_path: str) -> str:
        """从 Word(.docx) 文件中提取文本，包含段落与表格内容。"""
        try:
            import docx  # python-docx
        except ImportError:
            raise ValueError("缺少 python-docx 依赖，无法解析 Word 简历")
        try:
            document = docx.Document(docx_path)
            lines = [p.text for p in document.paragraphs if p.text and p.text.strip()]
            # 表格里常放教育/工作经历，一并提取
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            text = "\n".join(lines).strip()
            if not text:
                raise ValueError("Word 文档为空或无可提取文本")
            return text
        except Exception as e:
            logging.error(f"Word 文本提取失败: {e}")
            raise ValueError(f"Word 解析失败: {e}")
    
    def _call_llm(self, system_prompt: str, user_prompt: str) -> str:
        """调用统一LLM客户端"""
        return self.llm.chat(system_prompt, user_prompt, response_format=None, temperature=0.2)
    
    def extract_resume_info(self, resume_text: str) -> Dict[str, Any]:
        """使用LLM提取简历结构化信息"""
        system_prompt = (
            "你是一位专业的简历解析专家。请从简历文本中提取结构化信息。"
            "必须严格按照JSON格式输出，不要包含任何其他文字。"
        )
        
        user_prompt = (
            f"请从以下简历文本中提取信息，并以JSON格式输出：\n\n"
            f"{resume_text}\n\n"
            "输出格式：\n"
            "{\n"
            '  "name": "姓名",\n'
            '  "email": "邮箱",\n'
            '  "phone": "电话",\n'
            '  "education": [\n'
            '    {"school": "学校", "degree": "学位", "major": "专业", "year": "年份"}\n'
            '  ],\n'
            '  "experience": [\n'
            '    {"company": "公司", "position": "职位", "duration": "时间", "description": "描述"}\n'
            '  ]\n'
            "}\n"
            "如果某项信息不存在，使用空字符串或空数组。"
        )
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            # 尝试提取JSON
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                extracted = json.loads(json_match.group())
            else:
                # 如果找不到JSON，尝试直接解析
                extracted = json.loads(response)
            
            return {
                "name": extracted.get("name", ""),
                "email": extracted.get("email", ""),
                "phone": extracted.get("phone", ""),
                "education": extracted.get("education", []),
                "experience": extracted.get("experience", []),
            }
        except Exception as e:
            logging.error(f"简历信息提取失败: {e}")
            # 返回空结构
            return {
                "name": "",
                "email": "",
                "phone": "",
                "education": [],
                "experience": [],
            }
    
    def score_resume_skills(self, resume_text: str, extracted_info: Dict[str, Any]) -> List[Dict[str, Any]]:
        """为简历技能打分：先由LLM从所有 level_3rd 中选出<=10个适合用户的岗位类别，再基于这些类别汇总tags并打分"""
        # 构建简历文本
        profile_text = self._build_profile_text(resume_text, extracted_info)
        
        # 先选择适合的 level_3rd（最多10个）
        selected_level3 = self._select_level3_via_llm(profile_text, self.level3_list)
        logging.info(f"LLM 选择的 level_3rd: {selected_level3}")

        # 基于 level_3rd 汇总核心标签
        core_tags: List[str] = []
        for lv3 in selected_level3:
            core_tags.extend(self.level3_to_tags.get(lv3, []))
        # 兜底+常见标签
        default_common_tags = [
            "Python", "Java", "JavaScript", "C++", "SQL",
            "机器学习", "深度学习", "算法", "数据结构",
            "PyTorch", "TensorFlow", "Docker", "Kubernetes",
            "NLP", "计算机视觉", "推荐系统"
        ]

        # 结合文本直匹配的标签
        text_based_tags = self._select_candidate_tags(profile_text)

        # 汇总候选标签（核心优先），去重
        combined_tags_seq = core_tags + default_common_tags + text_based_tags
        seen = set()
        candidate_tags: List[str] = []
        for t in combined_tags_seq:
            if t and t not in seen and t in self.all_tags:
                seen.add(t)
                candidate_tags.append(t)
        # 控制数量
        candidate_tags = candidate_tags[:80]
        
        if not candidate_tags:
            logging.warning("未找到候选技能标签")
            return []
        
        # 调用LLM进行技能评分
        system_prompt = (
            "你是一位顶级的AI职业技能评估专家，以严格和精确著称。"
            "你的任务是基于用户简历，为【待评分标签列表】中的每一个标签进行1-5分的熟练度评估。\n\n"
            f"{COMMON_SCORING_RULES_V4}\n\n"
            "## 最终输出指令\n\n"
            "1. 格式: 你的回答只能包含 `标签:分数` 对，并用单个空格分隔。\n"
            "2. 示例: `Java:4 Spring Boot:3 MySQL:3`\n"
            "禁止任何其他文字、符号、换行、标题、解释。"
        )
        
        user_prompt = (
            f"{profile_text}\n\n"
            "### 任务详情\n\n"
            f"请为以下【待评分标签列表】中的每个标签打分。\n"
            f"【待评分标签列表】: {', '.join(candidate_tags[:50])}"  # 扩至最多50个
        )
        
        try:
            llm_reply = self._call_llm(system_prompt, user_prompt)
            scored_pairs = parse_llm_response(llm_reply, set(candidate_tags))
            
            # 转换为前端需要的格式
            # 排序：核心标签优先，其次分数降序、再按名称
            core_set = set(core_tags)
            scored_pairs_sorted = sorted(
                scored_pairs,
                key=lambda x: (
                    0 if x[0] in core_set else 1,
                    -x[1],
                    x[0]
                )
            )

            skills = []
            for tag_name, score in scored_pairs_sorted:
                # 确定分类
                category = self._categorize_skill(tag_name)
                skills.append({
                    "id": f"skill_{len(skills)}",
                    "resume_id": "temp",  # 会在API层替换
                    "skill_name": tag_name,
                    "score": score,
                    "category": category
                })
            
            return skills
        except Exception as e:
            logging.error(f"技能评分失败: {e}")
            return []

    def _select_level3_via_llm(self, profile_text: str, level3_all: List[str]) -> List[str]:
        """让LLM从所有 level_3rd 中选出 <=10 个最契合的岗位类别"""
        if not level3_all:
            return []
        level3_excerpt = ", ".join(level3_all[:200])  # 控制提示长度
        system_prompt = (
            "你是一位资深的职位画像分析专家。请从给定的岗位类别(level_3rd)列表中，"
            "选出最契合该候选人的最多10个类别。严格只返回逗号分隔的类别名称列表，不要包含其他任何符号或解释。"
        )
        user_prompt = (
            "### 候选人简历\n"
            f"{profile_text}\n\n"
            "### 可选的岗位类别(level_3rd)（不完整摘录，按语义选择即可）\n"
            f"{level3_excerpt}\n\n"
            "输出要求：只输出选中的类别，使用中文逗号或英文逗号分隔，最多10个。"
        )
        try:
            reply = self._call_llm(system_prompt, user_prompt)
            # 清洗成列表
            reply_clean = reply.replace("\n", " ").replace("，", ",")
            parts = [p.strip() for p in reply_clean.split(",") if p.strip()]
            # 只保留在完整列表中的合法项
            allowed = set(level3_all)
            selected = []
            for p in parts:
                if p in allowed and p not in selected:
                    selected.append(p)
                if len(selected) >= 10:
                    break
            # 如果为空，用简单文本匹配兜底
            if not selected:
                text_norm = profile_text.lower()
                for lv3 in level3_all:
                    if len(selected) >= 5:
                        break
                    if lv3 and lv3.lower() in text_norm:
                        selected.append(lv3)
            return selected[:10]
        except Exception as e:
            logging.warning(f"LLM 选择 level_3rd 失败，使用兜底逻辑: {e}")
            return level3_all[:5]
    
    def _build_profile_text(self, resume_text: str, extracted_info: Dict[str, Any]) -> str:
        """构建简历文本用于技能评估"""
        lines = [resume_text]
        
        if extracted_info.get("name"):
            lines.append(f"姓名: {extracted_info['name']}")
        if extracted_info.get("education"):
            lines.append("\n教育背景:")
            for edu in extracted_info["education"]:
                lines.append(f"- {edu.get('school', '')} {edu.get('degree', '')} {edu.get('major', '')}")
        if extracted_info.get("experience"):
            lines.append("\n工作经历:")
            for exp in extracted_info["experience"]:
                lines.append(f"- {exp.get('company', '')} {exp.get('position', '')}")
                lines.append(f"  {exp.get('description', '')}")
        
        return "\n".join(lines)
    
    def _select_candidate_tags(self, text: str) -> List[str]:
        """基于简历文本选择候选技能标签"""
        text_lower = text.lower()
        candidates = []
        
        # 直接匹配
        for tag in self.all_tags:
            if tag.lower() in text_lower:
                candidates.append(tag)
                if len(candidates) >= 50:  # 限制数量
                    break
        
        # 如果不够，添加一些常见标签
        if len(candidates) < 20:
            common_tags = [
                "Python", "Java", "JavaScript", "C++", "SQL",
                "机器学习", "深度学习", "算法", "数据结构",
                "React", "Vue", "Spring Boot", "Docker", "Kubernetes"
            ]
            for tag in common_tags:
                if tag in self.all_tags and tag not in candidates:
                    candidates.append(tag)
        
        return candidates[:50]  # 最多50个
    
    def _categorize_skill(self, skill_name: str) -> str:
        """对技能进行分类"""
        skill_lower = skill_name.lower()
        
        if any(kw in skill_lower for kw in ["python", "java", "javascript", "c++", "go", "rust"]):
            return "编程语言"
        elif any(kw in skill_lower for kw in ["react", "vue", "angular", "spring", "django", "框架"]):
            return "框架工具"
        elif any(kw in skill_lower for kw in ["mysql", "redis", "mongodb", "数据库"]):
            return "数据库"
        elif any(kw in skill_lower for kw in ["机器学习", "深度学习", "算法", "ai", "人工智能"]):
            return "核心能力"
        elif any(kw in skill_lower for kw in ["docker", "kubernetes", "git", "ci/cd", "工具"]):
            return "工具"
        else:
            return "其他"
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """解析简历的主方法（支持 PDF / Word .docx）"""
        logging.info(f"开始解析简历: {file_path}")

        # 1. 按文件类型提取文本
        resume_text = self.extract_text(file_path)
        logging.info(f"提取文本长度: {len(resume_text)} 字符")
        if not resume_text.strip():
            raise ValueError("简历内容为空，无法解析（可能是扫描件或加密文件）")
        
        # 2. 提取结构化信息
        extracted_info = self.extract_resume_info(resume_text)
        logging.info("简历信息提取完成")
        
        # 3. 技能评分
        skills = self.score_resume_skills(resume_text, extracted_info)
        logging.info(f"技能评分完成，共 {len(skills)} 个技能")
        
        from datetime import datetime
        return {
            "extracted_info": extracted_info,
            "skills": skills,
            "upload_date": datetime.now().isoformat(),
        }

