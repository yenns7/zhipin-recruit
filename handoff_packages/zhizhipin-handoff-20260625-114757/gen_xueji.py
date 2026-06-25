# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), '000000')
        borders.append(e)
    tblPr.append(borders)

doc = Document()

# 默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cell(cell, text, bold=False, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('广东技术师范大学学生学籍表')
r.font.name = '宋体'
r.font.size = Pt(16)
r.font.bold = True
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 顶部一行
top = doc.add_paragraph()
rr = top.add_run('学院：计算机科学学院    专业（大类）：电子商务（师范）    年级：2024级D班    学号：2024035424019')
rr.font.name = '宋体'; rr.font.size = Pt(10.5)
rr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ---- 基本信息表 ----
basic = [
    ('姓名', '李淑娟', '曾用名', '无'),
    ('性别', '女', '民族', '汉族'),
    ('政治面貌', '共青团员', '入团（党）时间', '2016年5月'),
    ('籍贯', '广东省梅州市', '出生地', '广东省梅州市'),
    ('生源所在地', '广东省梅州市', '入学时间', '2024年9月'),
    ('层次', '全日制专升本', '学制', '2'),
    ('出生日期', '2001年04月04日', '身份证号', '441424200104044848'),
    ('考生类别', '城镇往届', '入学前毕业学校或单位', '广东女子职业技术学院'),
    ('家长姓名', '李志达', '联系电话', '18602006269'),
]
t = doc.add_table(rows=len(basic) + 1, cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t)
for i, row in enumerate(basic):
    for j, val in enumerate(row):
        set_cell(t.rows[i].cells[j], val, bold=(j % 2 == 0), align='center')
# 末行：家庭通讯地址跨列
set_cell(t.rows[-1].cells[0], '家庭通讯地址', bold=True, align='center')
addr = t.rows[-1].cells[1].merge(t.rows[-1].cells[2]).merge(t.rows[-1].cells[3])
set_cell(addr, '广东省广州市白云区同和街同和西路51号A栋', align='center')

# ---- 学历及经历 ----
doc.add_paragraph()
h2 = doc.add_paragraph()
hr = h2.add_run('学历及经历')
hr.font.name = '宋体'; hr.font.size = Pt(12); hr.font.bold = True
hr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

edu_head = ('起止年月', '学校或单位名称', '任何职务')
edu = [
    ('2008.9 - 2014.6', '广东省梅州市远光小学', '语文科代表'),
    ('2014.9 - 2017.6', '广东省梅州市兴华中学', '数学科代表'),
    ('2017.9 - 2020.6', '广东省梅州市五华县高级中学', '英语科代表'),
    ('2020.9 - 2021.6', '广东省梅州市五华县高级中学', '无'),
    ('2021.9 - 2024.6', '广东女子职业技术学院', '宣传委员'),
    ('2024.9 - 2026.6', '广东技术师范大学', '宣传委员'),
]
te = doc.add_table(rows=len(edu) + 1, cols=3)
te.style = 'Table Grid'
te.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(te)
for j, val in enumerate(edu_head):
    set_cell(te.rows[0].cells[j], val, bold=True, align='center')
for i, row in enumerate(edu):
    for j, val in enumerate(row):
        set_cell(te.rows[i + 1].cells[j], val, align='center')

# ---- 家庭主要成员 ----
doc.add_paragraph()
h3 = doc.add_paragraph()
hr3 = h3.add_run('家庭主要成员及主要社会关系')
hr3.font.name = '宋体'; hr3.font.size = Pt(12); hr3.font.bold = True
hr3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

fam_head = ('姓名', '关系', '年龄', '政治面貌', '工作单位及职务')
fam = [
    ('李志达', '父亲', '50', '群众', '个体户'),
    ('李秋云', '母亲', '49', '群众', '个体户'),
    ('李娜', '姐妹', '20', '团员', '广东食品药品职业学院 学生'),
]
tf = doc.add_table(rows=len(fam) + 1, cols=5)
tf.style = 'Table Grid'
tf.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(tf)
for j, val in enumerate(fam_head):
    set_cell(tf.rows[0].cells[j], val, bold=True, align='center')
for i, row in enumerate(fam):
    for j, val in enumerate(row):
        set_cell(tf.rows[i + 1].cells[j], val, align='center')

# ---- 说明 ----
doc.add_paragraph()
for line in [
    '说明：',
    '1. 类别、层次：全日制普通本科、全日制普通专科、全日制专升本。',
    '2. 考生类别分为：城镇应届、城镇往届、农村应届、农村往届。',
]:
    p = doc.add_paragraph()
    rl = p.add_run(line)
    rl.font.name = '宋体'; rl.font.size = Pt(9)
    rl._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

out = r'C:\Users\Administrator\Desktop\hl\广东技术师范大学学生学籍表.docx'
doc.save(out)
print('saved:', out)
