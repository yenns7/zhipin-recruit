# -*- coding: utf-8 -*-
"""以原件 (1).docx 为模板，填入学生数据，100% 保留原排版/合并/边框/字体/照片栏/第二页。"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SRC = r'C:\Users\Administrator\Documents\xwechat_files\wxid_8aefzf6omhms22_50f8\msg\file\2026-06\广东技术师范大学学生学籍表(1).docx'
OUT = r'C:\Users\Administrator\Desktop\hl\广东技术师范大学学生学籍表_已填写.docx'

doc = Document(SRC)

def fill(cell, text, size=12, align=None):
    """写入单元格，保留原段落属性，统一宋体。"""
    text = str(text)
    p = cell.paragraphs[0]
    # 清掉占位文字（如 省 市 / 年 月 日），保留段落
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rPr.rFonts.set(qn('w:eastAsia'), '宋体') if rPr.rFonts is not None else run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if align is not None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = align

def setcell(table, r, c, text, size=12):
    fill(table.rows[r].cells[c], text, size=size)

# ===== 顶部抬头段落（标题下方那一行：学院/专业/年级/班级/学号）=====
HEADER = {
    '学院': '计算机科学学院',
    '专业（大类）': '电子商务（师范）',
    '年级': '2024级D班',
    '学号': '2024035424019',
}
for para in doc.paragraphs:
    t = para.text
    if '学院' in t and '专业' in t and '学号' in t:
        # 在每个字段标签后补值
        new = t
        new = new.replace('学院', '学院：计算机科学学院 ', 1)
        new = new.replace('年级', '年级：2024级D班 ', 1)
        new = new.replace('班\n级', '').replace('班级', '', 1)
        new = new.replace('学号', '学号：2024035424019', 1)
        # 专业（大类）
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        rn = para.add_run('学院：计算机科学学院    专业（大类）：电子商务（师范）    年级：2024级D班    学号：2024035424019')
        rn.font.name = '宋体'; rn.font.size = Pt(12)
        rn._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), '宋体')
        break

t0 = doc.tables[0]
# 基本信息（值列：c1 / c8 / c13）
setcell(t0, 0, 1, '李淑娟');           setcell(t0, 0, 8, '无');          setcell(t0, 0, 13, '女')
setcell(t0, 1, 1, '汉族');             setcell(t0, 1, 8, '共青团员');     setcell(t0, 1, 13, '2016年5月')
setcell(t0, 2, 1, '广东省梅州市');      setcell(t0, 2, 8, '广东省梅州市');  setcell(t0, 2, 13, '广东省梅州市')
setcell(t0, 3, 1, '2024年9月');        setcell(t0, 3, 8, '全日制专升本');  setcell(t0, 3, 13, '2年')
# 出生日期(c3) / 身份证号(c12)
setcell(t0, 4, 3, '2001年04月04日');   setcell(t0, 4, 12, '441424200104044848')
# 考生类别(c4) / 入学前毕业学校(c13)
setcell(t0, 5, 4, '城镇往届');         setcell(t0, 5, 13, '广东女子职业技术学院')
# 家长姓名(c4) / 联系电话(c10)
setcell(t0, 6, 4, '李志达');           setcell(t0, 6, 10, '18602006269')
# 家庭通讯地址(c7)
setcell(t0, 7, 7, '广东省广州市白云区同和街同和西路51号A栋')

# 学历及经历（R09 起，值列 c2=起止年月 c6=学校 c15=职务）
edu = [
    ('2008.9-2014.6', '广东省梅州市远光小学', '语文科代表'),
    ('2014.9-2017.6', '广东省梅州市兴华中学', '数学科代表'),
    ('2017.9-2020.6', '广东省梅州市五华县高级中学', '英语科代表'),
    ('2020.9-2021.6', '广东省梅州市五华县高级中学', ''),
    ('2021.9-2024.6', '广东女子职业技术学院', '宣传委员'),
    ('2024.9-2026.6', '广东技术师范大学', '宣传委员'),
]
for i, (a, b, c) in enumerate(edu):
    r = 9 + i
    setcell(t0, r, 2, a); setcell(t0, r, 6, b); setcell(t0, r, 15, c)

# 家庭主要成员（R18 起，值列 c2=姓名 c6=关系 c9=年龄 c12=政治面目 c14=工作单位及职务）
fam = [
    ('李志达', '父亲', '50', '群众', '个体户'),
    ('李秋云', '母亲', '49', '群众', '个体户'),
    ('李娜',   '姐妹', '20', '团员', '广东食品药品职业学院 学生'),
]
for i, (n, rel, age, pol, work) in enumerate(fam):
    r = 18 + i
    setcell(t0, r, 2, n); setcell(t0, r, 6, rel); setcell(t0, r, 9, age)
    setcell(t0, r, 12, pol); setcell(t0, r, 14, work)

doc.save(OUT)
print('saved:', OUT)
